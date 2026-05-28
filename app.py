import streamlit as st
import google.generativeai as genai
import PyPDF2
import os
import glob
from docx import Document
from io import BytesIO
import urllib.parse  # 💡 로컬 파일(HTML) 경로를 한글로 변환하고 읽기 위해 추가된 모듈

st.set_page_config(page_title="audskal의 학교생활기록부 분석", layout="wide")
st.title("🏫 객관적이고 체계적인 학생부 분석")
st.markdown("API 키에 맞는 최적의 AI 모델을 자동으로 찾아내어 생기부를 체계적으로 분석합니다.")

@st.cache_data(show_spinner=False)
def load_reference_pdfs(pdf_list):
    text = ""
    for pdf_file in pdf_list:
        with open(pdf_file, "rb") as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
    return text

with st.sidebar:
    st.header("🔑 기본 설정")
    api_key = st.text_input("API 키를 입력하세요", type="password")
    st.markdown("[👉 무료 API 키 발급받기](https://aistudio.google.com/app/apikey)")
    
    st.markdown("---")
    st.subheader("📚 내장된 평가 기준 파일 (참고용)")
    pdf_files = glob.glob("*.pdf")
    if pdf_files:
        for f in pdf_files:
            st.write(f"- {f}")
    else:
        st.error("폴더에 기준 PDF 파일이 없습니다!")

col1, col2 = st.columns([1, 1])

with col1:
    st.subheader("1. 학교생활기록부 데이터 입력")
    st.info("💡 나이스(NEIS) 원본 PDF는 보안상 안 읽히는 경우가 많습니다. 가급적 아래 빈칸에 내용을 직접 긁어서 붙여넣으세요!")
    
    student_file = st.file_uploader("📂 학생 생기부 파일 (PDF) 업로드", type=["pdf"])
    st.markdown("**-- 또는 --**")
    student_text_input = st.text_area("📝 생기부 내용 직접 붙여넣기 (추천)", height=250)

with col2:
    st.subheader("2. 학생부 분석을 위한 추가 정보 입력")
    teacher_context = st.text_area(
        "💡 특이사항 및 희망 전공 (예: 전기전자공학 진학 희망)", 
        height=70
    )
    
    # --- [수정된 부분] 선생님이 요청하신 로컬 링크를 기본값으로 셋팅 ---
    default_book_link = "file:///D:/%EC%9B%90%EB%93%9C%EB%9D%BC%EC%9D%B4%EB%B8%8C/OneDrive%20-%20%EC%9A%B8%EC%82%B0%EA%B4%91%EC%97%AD%EC%8B%9C%EA%B5%90%EC%9C%A1%EC%B2%AD/%EB%AC%B8%EC%84%9C/%EC%B9%B4%EC%B9%B4%EC%98%A4%ED%86%A1%20%EB%B0%9B%EC%9D%80%20%ED%8C%8C%EC%9D%BC/%EB%AF%B8%EB%9E%98%EB%A5%BC%20%EC%97%AC%EB%8A%94%20%EC%84%9C%EC%9E%AC(%EB%82%B4%EC%9D%BC%EA%B5%90%EC%9C%A1%20%EA%B6%8C%EC%9E%A5%EB%8F%84%EC%84%9C%20%EA%B4%80%EB%A0%A8%20%EA%B8%B0%EC%82%AC%20%EA%B2%80%EC%83%89-%EC%8B%A0%EC%84%A0%EC%97%AC%EA%B3%A0%20%EC%9E%84%EC%A2%85%EC%9A%B0).html"
    
    book_reference = st.text_input(
        "🔗 추천 도서 참고 링크 또는 목록 (선택)", 
        value=default_book_link,  # 화면을 켜면 이 값이 항상 기본으로 들어가게 됩니다.
        placeholder="도서 링크 또는 목록 입력 (비워두면 AI 자체 데이터 활용)"
    )
    
    submit_btn = st.button("↵ 🚀 심층 분석 시작 (클릭)", type="primary", use_container_width=True)

st.markdown("---")

def create_word_file(text):
    doc = Document()
    doc.add_heading('AI 생기부 분석 결과 보고서', 0)
    doc.add_paragraph(text)
    
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

if submit_btn:
    if not api_key:
        st.error("왼쪽에 API 키를 먼저 입력해 주세요!")
    elif not pdf_files:
        st.error("기준이 될 PDF 파일이 폴더에 없습니다!")
    elif not student_file and not student_text_input.strip():
        st.error("학생의 생기부 파일(PDF)을 업로드하거나 텍스트를 직접 붙여넣어 주세요!")
    else:
        status_box = st.empty()
        
        try:
            status_box.info("⏳ [진행상황 1/4] 내장된 가이드북(PDF)을 읽고 암기하는 중입니다...")
            reference_text = load_reference_pdfs(pdf_files)
            
            status_box.info("⏳ [진행상황 2/4] 학생의 생기부 데이터를 추출하는 중입니다...")
            student_data_text = ""
            
            if student_text_input.strip():
                student_data_text = student_text_input
            elif student_file:
                student_pdf_reader = PyPDF2.PdfReader(student_file)
                for page in student_pdf_reader.pages:
                    text = page.extract_text()
                    if text:
                        student_data_text += text + "\n"
            
            if not student_data_text.strip():
                raise Exception("업로드하신 PDF 파일에서 글씨를 읽을 수 없습니다! PDF 대신 왼쪽 빈칸에 생기부 내용을 직접 마우스로 긁어서 붙여넣어 주세요.")
            
            # --- 💡 [마법의 코드 추가] 로컬 HTML 파일을 파이썬이 대신 읽어서 AI에게 전달 ---
            status_box.info("📚 [도서 연동] 추천 도서 데이터를 불러오는 중입니다...")
            actual_book_data = book_reference
            
            # 링크가 file:/// 로 시작한다면 내 PC에 있는 파일이므로 열어서 읽어옵니다.
            if book_reference.startswith("file:///"):
                try:
                    # %EC%9B 와 같이 깨진 주소를 원래 한글 폴더명으로 복구합니다.
                    local_path = urllib.parse.unquote(book_reference.replace("file:///", ""))
                    if os.path.exists(local_path):
                        with open(local_path, "r", encoding="utf-8", errors='ignore') as f:
                            actual_book_data = f.read() # HTML 문서 내용을 통째로 변수에 저장
                    else:
                        st.warning("⚠️ 지정된 경로에서 도서 HTML 파일을 찾을 수 없어 AI 자체 데이터를 대신 활용합니다.")
                        actual_book_data = "제공된 파일 없음. AI 내장 권장 도서 활용 요망."
                except Exception as e:
                    actual_book_data = "제공된 파일 없음. AI 내장 권장 도서 활용 요망."

            status_box.warning("🔍 [진행상황 3/4] 최적의 구글 AI 모델을 탐색 중입니다...")
            genai.configure(api_key=api_key)
            
            best_model_name = ""
            for m in genai.list_models():
                if 'generateContent' in m.supported_generation_methods:
                    best_model_name = m.name.replace("models/", "")
                    if 'flash' in best_model_name or 'pro' in best_model_name:
                        break 
            
            if best_model_name == "":
                raise Exception("사용할 수 있는 AI 모델이 없습니다.")
            
            status_box.success(f"🤖 [진행상황 4/4] AI 엔진('{best_model_name}') 장착 완료! 객관적 경쟁력 분석을 시작합니다...")
            model = genai.GenerativeModel(best_model_name)
            
            prompt = f"""
            당신은 20년 경력의 대한민국 최고 수석 진학 상담 교사입니다.
            아래에 제공된 [대학 평가 기준 자료]는 훌륭한 학생부를 판단하기 위한 '참고용 범용 벤치마크(기준점)'입니다.
            이 기준점들에 비추어 보았을 때, [업로드된 학생의 생기부 내용]이 가진 객관적인 경쟁력과 역량 수준을 날카롭게 분석해 주세요.

            🚨 [주의: 특정 대학 편향 금지] 🚨
            - 특정 대학교에 지원한다는 가정하에 작성하지 마세요. 
            - 평가 기준 자료에 등장하는 특정 대학교 이름이나 '목표 대학'이라는 단어를 억지로 출력하지 마세요. 
            - 오직 '상위권 대학들이 공통으로 요구하는 역량'을 잣대로 삼아, 이 학생부 자체가 가진 경쟁력과 전공 적합성에만 집중하세요.

            🚨 [절대 엄수 - 팩트 체크 및 소설 작성 금지 규칙!] 🚨
            1. 팩트 기반 작성 (할루시네이션 절대 금지): 
               - 반드시 [업로드된 학생의 생기부 내용]에 '실제로 적혀있는 학년'과 '실제로 한 활동'만 가지고 분석하세요.
               - 생기부에 없는 내용, 학년, 과목명, 활동명은 단 한 글자도 지어내면 안 됩니다.
            2. 예시 내용 복사 금지: 
               - 아래의 [작성 예시]는 구조와 문체를 보여주기 위함입니다. 빈칸에 반드시 학생의 '실제 데이터'만 채워 넣으세요.
            3. 🆕 학년별 기록 부재를 약점으로 지적 절대 금지!: 
               - 본 분석 시점은 학기 진행 중이며, 특히 3학년의 경우 아직 학생부 기록이 입력되지 않은 것이 정상임.
               - "3학년 기록 부재" 등 학년별 기록의 부재를 약점으로 지적 절대 금지! 오직 '실제 입력된 활동 내용 자체의 질적 한계'에만 집중할 것.

            🚨 [형식 및 문체 규칙] 🚨
            1. 압축 서술: 사소한 활동은 버리고, 테마별로 가장 강력한 활동 단 2~3개만 엄선하여 3~4문장으로 압축할 것.
            2. 이중 출처 표기:
               - 문단 시작: 핵심 출처를 묶어서 `**[1학년 진로, 2학년 물리]** ` 형태로 표기.
               - 문장 끝: 해당 활동의 개별 출처를 `[1학년 진로]` 형태로 꼬리표 달기.
            3. 전 구간 개조식 어미 사용: 
               - 모든 문장의 끝은 '~함', '~임', '~됨', '~판단됨', '~요망됨' 으로 끝낼 것. ('~다', '~합니다' 절대 금지)

            [담당 교사의 특별 지시사항 및 희망 전공]
            {teacher_context if teacher_context else "특별한 지시사항 없음."}
            
            [추천 도서 참고 자료 (HTML 본문 또는 텍스트)]
            {actual_book_data}

            [대학 평가 기준 자료 (범용 벤치마크용)]
            {reference_text}

            [업로드된 학생의 생기부 내용 (100% 팩트)]
            {student_data_text}

            위의 규칙을 완벽히 지켜서, 학생의 실제 데이터만을 바탕으로 아래 5가지 양식에 맞추어 답변해 주세요.
            ### 1. 전공 적합성 및 주요 경쟁력
            ### 2. 범용 평가 기준에 비추어 볼 때 보완이 필요한 약점
            ### 3. 추천 심화 탐구 주제 및 면접 예상 질문 3가지
            ### 4. 종합 의견 및 향후 발전 방향
            ### 5. 맞춤형 추천 도서 및 연계 활동 제안 (반드시 [추천 도서 참고 자료]의 데이터를 우선적으로 분석 및 반영하여 희망 전공과 직결된 도서 3권을 추천하고, 각 도서를 읽은 후 세특에 녹여낼 수 있는 구체적인 '후속 탐구 활동'을 제안할 것)
            """
            
            response = model.generate_content(prompt)
            
            status_box.success("✅ [분석 완료!] 초고속 심층 분석이 완료되었습니다. 결과물을 확인해 주세요!")
            st.write(response.text)
            
            word_file = create_word_file(response.text)
            st.download_button(
                label="📥 분석 결과 워드(Word) 파일로 다운로드",
                data=word_file,
                file_name="생기부_분석결과_도서추천포함.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
        except Exception as e:
            status_box.error(f"오류가 발생했습니다: {e}")

# ===== 푸터(만든이 정보) =====
st.divider()
st.markdown("""
<div style='text-align: center; color: gray; padding: 20px; font-size: 13px;'>
    🏫 학교생활기록부 분석 시스템 v4.1 (도서 DB 자동연동 기술 적용)<br>
    만든이: <b>신선여자고등학교 김명남</b><br>
    🗓️ 2026.04
</div>
""", unsafe_allow_html=True)
